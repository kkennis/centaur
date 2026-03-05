"""Closing document generators — Board Consent, Stockholder Consent, MRL, certificates.

Generates properly formatted DOCX documents using python-docx (no template files).
Each generator takes deal parameters from a TermSheet and produces bytes.
"""

import io
from datetime import UTC, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .docx_pipeline import format_money
from .models import BoardRights, TermSheet


def _styled_doc() -> Document:
    """Create a blank Document with Paradigm-standard font defaults."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    return doc


def _add_title(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.bold = True
        run.font.size = Pt(12)
        if i < len(lines) - 1:
            para.add_run("\n")


def _add_body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def _add_resolved(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run_label = para.add_run("RESOLVED")
    run_label.bold = True
    run_label.font.size = Pt(11)
    para.add_run(f", {text}")


def _add_further_resolved(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run_label = para.add_run("FURTHER RESOLVED")
    run_label.bold = True
    run_label.font.size = Pt(11)
    para.add_run(f", {text}")


def _add_signature_block(doc: Document, name: str = "", title: str = "") -> None:
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("_" * 40)
    if name:
        doc.add_paragraph(f"Name: {name}")
    if title:
        doc.add_paragraph(f"Title: {title}")


def _effective_date(date: str | None) -> str:
    if date:
        return date
    return datetime.now(UTC).strftime("%B %d, %Y")


def generate_board_consent_docx(
    ts: TermSheet,
    effective_date: str | None = None,
    additional_resolutions: list[str] | None = None,
) -> bytes:
    """Generate a Board Consent approving a preferred stock financing."""
    doc = _styled_doc()
    company = ts.company_name
    series = ts.effective_series
    amt = format_money(ts.investment_amount)
    date_str = _effective_date(effective_date)

    _add_title(doc, [
        "ACTION BY UNANIMOUS WRITTEN CONSENT",
        f"OF THE BOARD OF DIRECTORS OF {company.upper()}",
        f"(a Delaware corporation)",
    ])

    doc.add_paragraph(f"Effective Date: {date_str}")
    doc.add_paragraph()

    _add_body(doc, (
        f'The undersigned, constituting all of the members of the Board of Directors '
        f'(the "Board") of {company}, a Delaware corporation (the "Company"), '
        f'hereby adopt the following resolutions by unanimous written consent '
        f'without a meeting pursuant to Section 141(f) of the Delaware General '
        f'Corporation Law and the Company\'s Bylaws:'
    ))

    _add_body(doc, (
        f'WHEREAS, the Board has determined that it is in the best interests of the '
        f'Company and its stockholders to effect a financing (the "Financing") through the '
        f'issuance and sale of shares of Series {series} Preferred Stock of the Company '
        f'for aggregate gross proceeds of up to {amt};'
    ))

    _add_body(doc, (
        f'WHEREAS, in connection with the Financing, the Board desires to authorize '
        f'the execution and delivery of: (i) that certain Amended and Restated Certificate '
        f'of Incorporation (the "Restated Certificate"), (ii) that certain Series {series} '
        f'Preferred Stock Purchase Agreement (the "Purchase Agreement"), (iii) that certain '
        f'Investors\' Rights Agreement, (iv) that certain Voting Agreement, (v) that certain '
        f'Right of First Refusal and Co-Sale Agreement, and (vi) such other agreements '
        f'and documents as may be required in connection therewith (collectively with the '
        f'Restated Certificate and the Purchase Agreement, the "Transaction Documents");'
    ))

    _add_body(doc, "NOW, THEREFORE, BE IT:")

    _add_resolved(doc, (
        f'that the Financing on the terms and conditions set forth in the Transaction '
        f'Documents is hereby approved and adopted in all respects.'
    ))

    _add_further_resolved(doc, (
        f'that the Restated Certificate, in substantially the form presented to the Board, '
        f'is hereby approved and adopted, and the officers of the Company are hereby '
        f'authorized and directed to file the Restated Certificate with the Secretary of '
        f'State of the State of Delaware.'
    ))

    _add_further_resolved(doc, (
        f'that the issuance and sale of up to such number of shares of Series {series} '
        f'Preferred Stock as set forth in the Purchase Agreement, at the price per share '
        f'set forth therein, is hereby approved.'
    ))

    _add_further_resolved(doc, (
        'that each of the Transaction Documents, in substantially the forms presented '
        'to the Board, is hereby approved, and the officers of the Company are hereby '
        'authorized and directed to execute and deliver each such document on behalf of '
        'the Company, with such changes therein as the officer executing the same shall '
        'approve, such approval to be conclusively evidenced by such execution.'
    ))

    if ts.option_pool_percent > 0:
        _add_further_resolved(doc, (
            f'that the Company\'s equity incentive plan is hereby amended to increase the '
            f'number of shares of Common Stock reserved for issuance thereunder such that, '
            f'following the Financing, the shares reserved for issuance under the plan '
            f'(including shares subject to outstanding awards) shall equal '
            f'{ts.option_pool_percent:g}% of the fully-diluted capitalization of the Company '
            f'on a post-Financing basis.'
        ))

    _add_further_resolved(doc, (
        'that the officers of the Company are hereby authorized and directed to take '
        'all such further actions and to execute and deliver all such further agreements, '
        'instruments, documents, and certificates as such officers shall deem necessary '
        'or advisable to carry out the purposes and intent of the foregoing resolutions, '
        'the execution thereof to be conclusive evidence of the Board\'s approval.'
    ))

    for resolution in additional_resolutions or []:
        _add_further_resolved(doc, resolution)

    _add_further_resolved(doc, (
        'that any and all actions heretofore taken by any officer or director of the '
        'Company in connection with the matters contemplated by the foregoing resolutions '
        'are hereby ratified, confirmed, approved and adopted as acts and deeds of the Company.'
    ))

    doc.add_paragraph()
    _add_body(doc, (
        'This action by unanimous written consent may be executed in one or more '
        'counterparts, each of which shall be deemed an original and all of which '
        'together shall constitute one and the same instrument. This consent shall '
        'be filed with the minutes of the proceedings of the Board.'
    ))

    _add_body(doc, "[Signature Page Follows]")

    doc.add_page_break()
    _add_body(doc, (
        "IN WITNESS WHEREOF, the undersigned have executed this Action by Unanimous "
        "Written Consent as of the date first written above."
    ))
    _add_signature_block(doc, title="Director")
    _add_signature_block(doc, title="Director")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_stockholder_consent_docx(
    ts: TermSheet,
    effective_date: str | None = None,
) -> bytes:
    """Generate a Stockholder Consent approving a preferred stock financing."""
    doc = _styled_doc()
    company = ts.company_name
    series = ts.effective_series
    date_str = _effective_date(effective_date)

    _add_title(doc, [
        "ACTION BY WRITTEN CONSENT",
        f"OF THE STOCKHOLDERS OF {company.upper()}",
        f"(a Delaware corporation)",
    ])

    doc.add_paragraph(f"Effective Date: {date_str}")
    doc.add_paragraph()

    _add_body(doc, (
        f'The undersigned, constituting the holders of a majority of the outstanding shares '
        f'of capital stock of {company}, a Delaware corporation (the "Company"), entitled to '
        f'vote on the matters set forth herein, hereby adopt the following resolutions by '
        f'written consent without a meeting pursuant to Section 228 of the Delaware General '
        f'Corporation Law and the Company\'s Bylaws:'
    ))

    _add_resolved(doc, (
        f'that the Amended and Restated Certificate of Incorporation of the Company '
        f'(the "Restated Certificate"), in substantially the form attached hereto as '
        f'Exhibit A, which authorizes the issuance of shares of Series {series} Preferred '
        f'Stock on the terms and conditions set forth therein, is hereby approved and '
        f'adopted in all respects.'
    ))

    _add_further_resolved(doc, (
        f'that the issuance and sale of shares of Series {series} Preferred Stock pursuant '
        f'to the terms of that certain Stock Purchase Agreement by and among the Company '
        f'and the investors listed therein (the "Purchase Agreement") is hereby approved.'
    ))

    _add_further_resolved(doc, (
        'that the execution and delivery by the Company of each of the Investors\' Rights '
        'Agreement, Voting Agreement, and Right of First Refusal and Co-Sale Agreement, '
        'each in substantially the forms presented to the stockholders, is hereby approved.'
    ))

    if ts.option_pool_percent > 0:
        _add_further_resolved(doc, (
            f'that the amendment to the Company\'s equity incentive plan to increase the '
            f'share reserve to {ts.option_pool_percent:g}% of the post-Financing fully-diluted '
            f'capitalization is hereby approved.'
        ))

    _add_further_resolved(doc, (
        'that the officers of the Company are hereby authorized and directed to take all '
        'such further actions as may be necessary or advisable to carry out the purposes '
        'and intent of the foregoing resolutions.'
    ))

    doc.add_paragraph()
    _add_body(doc, (
        'This action by written consent may be executed in one or more counterparts. '
        'Pursuant to Section 228(e) of the DGCL, prompt notice of the taking of this '
        'action shall be given to those stockholders who did not consent in writing hereto.'
    ))

    _add_body(doc, "[Signature Page Follows]")

    doc.add_page_break()
    _add_body(doc, (
        "IN WITNESS WHEREOF, the undersigned have executed this Action by Written "
        "Consent as of the date first written above."
    ))
    _add_signature_block(doc, title="Stockholder")
    _add_signature_block(doc, title="Stockholder")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_mrl_docx(
    ts: TermSheet,
    fund_name: str = "Paradigm Fund LP",
    effective_date: str | None = None,
) -> bytes:
    """Generate a Management Rights Letter for ERISA/VCOC qualification."""
    doc = _styled_doc()
    company = ts.company_name
    date_str = _effective_date(effective_date)

    para = doc.add_paragraph()
    para.add_run(date_str)
    doc.add_paragraph()
    _add_body(doc, f"{fund_name}")
    doc.add_paragraph()

    _add_body(doc, f"Re: Management Rights — {company}")
    doc.add_paragraph()

    _add_body(doc, f"Ladies and Gentlemen:")
    doc.add_paragraph()

    _add_body(doc, (
        f'This letter (this "Letter Agreement") is being delivered to {fund_name} '
        f'(the "Investor") in connection with the Investor\'s investment in {company}, '
        f'a Delaware corporation (the "Company").'
    ))

    _add_body(doc, (
        "The Company hereby agrees that the Investor shall be entitled to the following "
        "contractual management rights:"
    ))

    _add_body(doc, (
        '(a) The Investor shall have the right to consult with and advise management of the '
        'Company on significant business issues, including management\'s proposed annual '
        'operating plans, and management will meet with the Investor regularly during each '
        'year at the Company\'s facilities at mutually agreeable times for such consultation.'
    ))

    _add_body(doc, (
        '(b) The Investor shall have the right to examine the books and records of the '
        'Company and inspect its facilities and shall have the right to receive such '
        'financial and operating information as is provided to the Board of Directors of '
        'the Company. The Company shall provide reasonable access to the Investor for the '
        'purposes described herein.'
    ))

    has_board_seat = ts.board_rights in {BoardRights.SEAT, BoardRights.SEAT_AND_OBSERVER}
    has_observer = ts.board_rights in {BoardRights.OBSERVER, BoardRights.SEAT_AND_OBSERVER}

    if has_observer and not has_board_seat:
        _add_body(doc, (
            '(c) In the event that the Investor does not have a representative serving on '
            'the Company\'s Board of Directors, the Company shall invite a representative '
            'designated by the Investor to attend all meetings of the Board of Directors in '
            'a nonvoting observer capacity and shall provide such representative copies of '
            'all notices, minutes, consents, and other materials provided to the directors.'
        ))

    _add_body(doc, (
        "The rights described herein are intended to be contractual management rights "
        "within the meaning of 29 C.F.R. § 2510.3-101(d)(3)(ii) as required in order "
        "for the Investor to qualify as a venture capital operating company as defined in "
        "29 C.F.R. § 2510.3-101(d)."
    ))

    _add_body(doc, (
        "The rights described herein shall terminate upon the earlier of (i) the "
        "consummation of an initial public offering of the Company's securities or "
        "(ii) the date on which the Investor ceases to hold any securities of the Company."
    ))

    doc.add_paragraph()
    _add_body(doc, "Very truly yours,")
    doc.add_paragraph()
    _add_body(doc, company.upper())
    _add_signature_block(doc, title="Authorized Signatory")
    doc.add_paragraph()

    _add_body(doc, "ACKNOWLEDGED AND AGREED:")
    doc.add_paragraph()
    _add_body(doc, fund_name.upper())
    _add_signature_block(doc, title="Authorized Signatory")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_secretary_certificate_docx(
    ts: TermSheet,
    secretary_name: str = "",
    effective_date: str | None = None,
) -> bytes:
    """Generate a Secretary's Certificate for a financing closing."""
    doc = _styled_doc()
    company = ts.company_name
    series = ts.effective_series
    date_str = _effective_date(effective_date)

    _add_title(doc, [
        "SECRETARY'S CERTIFICATE",
        f"OF {company.upper()}",
    ])

    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph()

    name_str = secretary_name or "[Secretary Name]"
    _add_body(doc, (
        f'I, {name_str}, hereby certify that I am the duly elected, qualified, and acting '
        f'Secretary of {company}, a Delaware corporation (the "Company"), and that as such '
        f'I am authorized to execute this certificate on behalf of the Company in connection '
        f'with the closing of the Company\'s Series {series} Preferred Stock financing '
        f'(the "Financing").'
    ))

    _add_body(doc, "I further certify that:")

    _add_body(doc, (
        f'1. Attached hereto as Exhibit A is a true and correct copy of the Amended and '
        f'Restated Certificate of Incorporation of the Company, as filed with the Secretary '
        f'of State of the State of Delaware, which is in full force and effect as of the date '
        f'hereof.'
    ))

    _add_body(doc, (
        '2. Attached hereto as Exhibit B is a true and correct copy of the Bylaws of the '
        'Company, as currently in effect.'
    ))

    _add_body(doc, (
        '3. Attached hereto as Exhibit C is a true and correct copy of the resolutions of '
        'the Board of Directors of the Company authorizing the Financing and the transactions '
        'contemplated thereby, which resolutions have not been amended, modified, or rescinded '
        'and remain in full force and effect.'
    ))

    _add_body(doc, (
        '4. The following persons are the duly elected or appointed officers of the Company, '
        'holding the offices set forth opposite their respective names, and the signatures '
        'set forth opposite their names are their genuine signatures:'
    ))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, header in enumerate(["Name", "Title", "Signature"]):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for _ in range(3):
        table.add_row()

    doc.add_paragraph()
    _add_body(doc, "IN WITNESS WHEREOF, I have executed this certificate as of the date first written above.")
    _add_signature_block(doc, name=name_str, title="Secretary")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_compliance_certificate_docx(
    ts: TermSheet,
    officer_name: str = "",
    officer_title: str = "Chief Executive Officer",
    effective_date: str | None = None,
) -> bytes:
    """Generate an Officer's Compliance Certificate for a financing closing."""
    doc = _styled_doc()
    company = ts.company_name
    series = ts.effective_series
    date_str = _effective_date(effective_date)

    _add_title(doc, [
        "OFFICER'S COMPLIANCE CERTIFICATE",
        f"OF {company.upper()}",
    ])

    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph()

    name_str = officer_name or "[Officer Name]"
    _add_body(doc, (
        f'I, {name_str}, {officer_title} of {company}, a Delaware corporation '
        f'(the "Company"), hereby certify, in connection with the closing of the '
        f'Company\'s Series {series} Preferred Stock financing pursuant to that certain '
        f'Stock Purchase Agreement, dated as of {date_str} (the "Purchase Agreement"), '
        f'by and among the Company and the investors listed therein, that:'
    ))

    _add_body(doc, (
        '1. The representations and warranties of the Company contained in the Purchase '
        'Agreement are true and correct in all material respects as of the date hereof.'
    ))

    _add_body(doc, (
        '2. All covenants, agreements, and conditions contained in the Purchase Agreement '
        'to be performed or complied with by the Company on or before the date hereof have '
        'been performed or complied with in all material respects.'
    ))

    _add_body(doc, (
        '3. No officer, director, or twenty percent (20%) stockholder of the Company is '
        'subject to any of the "bad actor" disqualifying events described in Rule 506(d)(1)(i) '
        'through (viii) under the Securities Act of 1933, as amended.'
    ))

    _add_body(doc, (
        '4. The Company has filed or caused to be filed the Amended and Restated Certificate '
        'of Incorporation with the Secretary of State of the State of Delaware.'
    ))

    doc.add_paragraph()
    _add_body(doc, "IN WITNESS WHEREOF, I have executed this certificate as of the date first written above.")
    _add_signature_block(doc, name=name_str, title=officer_title)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
